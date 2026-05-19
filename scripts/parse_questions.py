"""Parse the CCNA Cert Empire .docx into the quiz DB.

The source is `docs/200-301 Cert Empire x990.docx` — a Word doc where each
question lives in its own 2-column table. The correct choice is marked by
light-green run shading (fill = ``e6f0e6``). Images embedded in a question
table are extracted to ``data/quiz_images/`` and referenced by filename in
the ``questions.image_filenames`` column.

Usage from the project root:

    python scripts/parse_questions.py --probe   # summary only, no DB writes
    python scripts/parse_questions.py           # full parse + image extract

Re-running is safe — the parser deletes and re-inserts every row, and
images are re-extracted by question id.
"""

from __future__ import annotations

import argparse
import asyncio
import json
import logging
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import aiosqlite
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from core.exam_pools import POOL_VALUES, TOPIC_TO_POOL  # noqa: E402

SOURCE_DOCX = ROOT / "docs" / "200-301 Cert Empire x990.docx"
IMAGES_DIR  = ROOT / "data" / "quiz_images"
DB_PATH     = ROOT / "database" / "labs.db"

GREEN_FILL  = "e6f0e6"

W_SHD       = qn("w:shd")
W_FILL      = qn("w:fill")
W_RPR       = qn("w:rPr")
A_NS        = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS        = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
BLIP_TAG    = f"{{{A_NS}}}blip"
EMBED_ATTR  = f"{{{R_NS}}}embed"

ALLOWED_IMAGE_EXTS = frozenset({"png", "jpg", "gif", "webp", "bmp"})

HEADER_RE = re.compile(
    r"Question\s*#\s*(\d+).*?Topic\s+(\d+)\s*-\s*Exam\s+Pool\s+([A-Z])",
    re.DOTALL,
)
THAI_RE = re.compile(r"[฀-๿]")

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedQuestion:
    id:              int
    pool:            str
    topic:           int
    prompt_en:       str
    prompt_th:       str | None
    choices:         list[dict[str, str]]
    correct_labels:  list[str]
    explanation:     str | None
    source_table:    int
    image_filenames: list[str] = field(default_factory=list)
    needs_review:    bool = False


def parse_header_text(text: str) -> tuple[int, int, str] | None:
    """Extract (question_id, topic, pool) from the joined row-0 cells.

    Returns None if the text doesn't look like a question header.
    """
    m = HEADER_RE.search(text)
    if not m:
        return None
    qid = int(m.group(1))
    topic = int(m.group(2))
    pool = m.group(3)
    if topic not in TOPIC_TO_POOL or pool not in POOL_VALUES:
        return None
    return qid, topic, pool


def is_mostly_thai(text: str) -> bool:
    """True when Thai characters outnumber ASCII letters in *text*."""
    if not text:
        return False
    thai = len(THAI_RE.findall(text))
    ascii_letters = sum(c.isascii() and c.isalpha() for c in text)
    return thai > ascii_letters


def cell_is_green(cell: _Cell) -> bool:
    """True if any run inside *cell* carries the green shading fill."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            rpr = run._r.find(W_RPR)
            if rpr is None:
                continue
            shd = rpr.find(W_SHD)
            if shd is None:
                continue
            fill = shd.get(W_FILL)
            if fill and fill.lower() == GREEN_FILL:
                return True
    return False


def _cell_text(cell: _Cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


def classify_table(table: Table, source_index: int) -> ParsedQuestion | None:
    """Turn a question table into a structured `ParsedQuestion`.

    Returns None if the table is not a recognisable question (no header).
    The classifier walks rows after the header and partitions cells into:

    * **duplicate rows** — both columns hold identical non-empty text
      → these are either the EN/TH prompt or the explanation
    * **choice rows** — the two columns differ → each non-empty cell is a
      choice; a green-shaded choice is marked correct
    * **empty rows** — skipped (typically the row reserved for an image)
    """
    rows = table.rows
    if not rows:
        return None

    header_cells = rows[0].cells
    if len(header_cells) < 2:
        return None
    header_text = f"{_cell_text(header_cells[0])} {_cell_text(header_cells[1])}"
    header = parse_header_text(header_text)
    if header is None:
        return None
    qid, topic, pool = header

    duplicate_rows: list[tuple[int, str]] = []
    choice_rows: list[tuple[int, list[tuple[str, bool]]]] = []
    for ri in range(1, len(rows)):
        cells = rows[ri].cells
        col_data = [(_cell_text(c), cell_is_green(c)) for c in cells]
        texts = [t for t, _ in col_data]
        if all(not t for t in texts):
            continue
        non_empty = {t for t in texts if t}
        if len(non_empty) == 1:
            duplicate_rows.append((ri, next(iter(non_empty))))
        else:
            choice_rows.append((ri, col_data))

    first_choice_ri = choice_rows[0][0] if choice_rows else float("inf")
    last_choice_ri = choice_rows[-1][0] if choice_rows else -1

    pre_choice_dups  = [t for ri, t in duplicate_rows if ri < first_choice_ri]
    post_choice_dups = [t for ri, t in duplicate_rows if ri > last_choice_ri]

    prompt_en = next((t for t in pre_choice_dups if not is_mostly_thai(t)), "")
    prompt_th = next((t for t in pre_choice_dups if is_mostly_thai(t)), None)

    choices: list[dict[str, str]] = []
    correct_labels: list[str] = []
    for _ri, col_data in choice_rows:
        for text, is_green in col_data:
            if not text:
                continue
            label = chr(ord("A") + len(choices))
            choices.append({"label": label, "text": text})
            if is_green:
                correct_labels.append(label)

    explanation = " ".join(t for t in post_choice_dups if t) or None

    review_reasons = []
    if not prompt_en:
        review_reasons.append("missing EN prompt")
    if len(choices) < 2:
        review_reasons.append(f"only {len(choices)} choice(s)")
    if not correct_labels:
        review_reasons.append("no green-marked answer")
    needs_review = bool(review_reasons)
    if needs_review:
        log.warning(
            "Q%04d (table %d): flagged — %s",
            qid, source_index, "; ".join(review_reasons),
        )

    return ParsedQuestion(
        id=qid,
        pool=pool,
        topic=topic,
        prompt_en=prompt_en,
        prompt_th=prompt_th,
        choices=choices,
        correct_labels=correct_labels,
        explanation=explanation,
        source_table=source_index,
        image_filenames=[],
        needs_review=needs_review,
    )


def extract_table_images(
    document: Document,
    table: Table,
    dest_dir: Path,
    qid: int,
) -> list[str]:
    """Extract every embedded image inside *table* to *dest_dir*.

    Returns a list of filenames (relative to *dest_dir*). Filenames follow
    ``Q-{id:04d}-{i}.{ext}``; ``i`` increments per image within the table.
    """
    filenames: list[str] = []
    blips = list(table._tbl.iter(BLIP_TAG))
    if not blips:
        return filenames

    dest_dir.mkdir(parents=True, exist_ok=True)
    for blip in blips:
        rid = blip.get(EMBED_ATTR)
        if not rid:
            continue
        try:
            image_part = document.part.related_parts[rid]
            ext = image_part.content_type.rsplit("/", 1)[-1].lower()
            if ext == "jpeg":
                ext = "jpg"
            if ext not in ALLOWED_IMAGE_EXTS:
                log.warning(
                    "Q%04d: unexpected image content-type %r — skipping",
                    qid, image_part.content_type,
                )
                continue
            fname = f"Q-{qid:04d}-{len(filenames)}.{ext}"
            (dest_dir / fname).write_bytes(image_part.blob)
            filenames.append(fname)
        except (KeyError, OSError, ValueError) as exc:
            log.warning("Q%04d: failed to extract image (rId=%s): %s", qid, rid, exc)
            continue
    return filenames


async def upsert_question(db: aiosqlite.Connection, q: ParsedQuestion) -> None:
    await db.execute(
        """
        INSERT OR REPLACE INTO questions
            (id, pool, topic, prompt_en, prompt_th, choices_json,
             correct_labels, explanation, source_table,
             image_filenames, needs_review)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            q.id,
            q.pool,
            q.topic,
            q.prompt_en,
            q.prompt_th,
            json.dumps(q.choices, ensure_ascii=False),
            json.dumps(q.correct_labels),
            q.explanation,
            q.source_table,
            json.dumps(q.image_filenames),
            int(q.needs_review),
        ),
    )


def probe(document: Document) -> int:
    """Print a structural summary without writing anything."""
    tables = document.tables
    total = len(tables)
    parsed = drag_drop = no_choices = no_correct = with_images = 0
    pool_counts: dict[str, int] = {p: 0 for p in POOL_VALUES}
    for i, table in enumerate(tables):
        q = classify_table(table, source_index=i)
        if q is None:
            continue
        parsed += 1
        pool_counts[q.pool] = pool_counts.get(q.pool, 0) + 1
        if len(q.choices) < 2:
            no_choices += 1
        if not q.correct_labels:
            no_correct += 1
        if list(table._tbl.iter(BLIP_TAG)):
            with_images += 1
        if "drag" in q.prompt_en.lower() and "drop" in q.prompt_en.lower():
            drag_drop += 1

    log.info("Tables in doc:        %d", total)
    log.info("Recognised questions: %d", parsed)
    log.info("Pool counts:          %s", pool_counts)
    log.info("With image(s):        %d", with_images)
    log.info("Drag-and-drop:        %d (will be flagged)", drag_drop)
    log.info("Missing choices:      %d", no_choices)
    log.info("No correct marker:    %d", no_correct)
    return 0


async def parse_all(document: Document) -> int:
    """Full parse: extract images, write rows, return exit code."""
    if not DB_PATH.exists():
        log.error("DB not found at %s — run the server once to init it.", DB_PATH)
        return 1

    parsed_total = with_images_total = flagged_total = 0
    seen_ids: set[int] = set()
    duplicate_ids: list[int] = []
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute("DELETE FROM questions")
        for i, table in enumerate(document.tables):
            q = classify_table(table, source_index=i)
            if q is None:
                continue
            if q.id in seen_ids:
                duplicate_ids.append(q.id)
            seen_ids.add(q.id)
            images = extract_table_images(document, table, IMAGES_DIR, q.id)
            if images:
                with_images_total += 1
            q = ParsedQuestion(
                id=q.id,
                pool=q.pool,
                topic=q.topic,
                prompt_en=q.prompt_en,
                prompt_th=q.prompt_th,
                choices=q.choices,
                correct_labels=q.correct_labels,
                explanation=q.explanation,
                source_table=q.source_table,
                image_filenames=images,
                needs_review=q.needs_review,
            )
            await upsert_question(db, q)
            parsed_total += 1
            if q.needs_review:
                flagged_total += 1
        await db.commit()

    log.info("Parsed:        %d", parsed_total)
    log.info("With images:   %d", with_images_total)
    log.info("Flagged:       %d", flagged_total)
    if duplicate_ids:
        log.warning(
            "Duplicate question ids overwritten (later wins): %s",
            sorted(set(duplicate_ids))[:20],
        )
    log.info("Done.")
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--probe", action="store_true",
                        help="Summarise structure without writing to DB.")
    args = parser.parse_args(argv)

    if not SOURCE_DOCX.exists():
        log.error("Source DOCX not found at %s", SOURCE_DOCX)
        return 1

    log.info("Loading %s ...", SOURCE_DOCX.name)
    try:
        document = Document(str(SOURCE_DOCX))
    except Exception as exc:  # noqa: BLE001 — surface any docx-load failure cleanly
        log.error("Failed to open DOCX: %s", exc)
        return 1

    if args.probe:
        return probe(document)
    return asyncio.run(parse_all(document))


if __name__ == "__main__":
    sys.exit(main())
