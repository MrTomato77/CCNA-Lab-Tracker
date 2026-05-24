"""Tests for the DOCX question parser.

Pure-function tests run anywhere. Integration tests against the real
Cert Empire docx are skipped if the (gitignored) source file is absent.
"""
from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table, _Cell

from scripts.parse_questions import (
    GREEN_FILL,
    SOURCE_DOCX,
    _cell_text,
    classify_table,
    extract_table_images,
    is_mostly_thai,
    parse_header_text,
)


# ---------- pure-function tests (no docx needed) ----------

def test_parse_header_text_extracts_id_topic_pool():
    text = "Question #42 Topic 3 - Exam Pool C"
    assert parse_header_text(text) == (42, 3, "C")


def test_parse_header_text_tolerates_extra_whitespace_and_split_lines():
    text = "Question #1\n   Topic 1 - Exam Pool A\n"
    assert parse_header_text(text) == (1, 1, "A")


def test_parse_header_text_returns_none_on_garbage():
    assert parse_header_text("Hello world") is None
    assert parse_header_text("") is None


def test_parse_header_text_rejects_unknown_topic_or_pool():
    assert parse_header_text("Question #1 Topic 9 - Exam Pool A") is None
    assert parse_header_text("Question #1 Topic 1 - Exam Pool Z") is None


def test_is_mostly_thai_detects_thai_dominant():
    assert is_mostly_thai("คำสั่งใด ที่ป้อนบนสวิตช์") is True


def test_is_mostly_thai_rejects_english_dominant():
    assert is_mostly_thai("Which command entered on a switch") is False


def test_is_mostly_thai_on_empty_string():
    assert is_mostly_thai("") is False


# ---------- integration tests (need real docx) ----------

needs_docx = pytest.mark.skipif(
    not SOURCE_DOCX.exists(),
    reason=f"source docx not available at {SOURCE_DOCX}",
)


@pytest.fixture(scope="module")
def document():
    from docx import Document
    return Document(str(SOURCE_DOCX))


@needs_docx
def test_classify_table_extracts_q1_ground_truth(document):
    """Q#1 has 4 choices A-D; the correct answer is D (forward-time 20),
    marked by the green run shading in the source doc."""
    q = classify_table(document.tables[0], source_index=0)

    assert q is not None
    assert q.id == 1
    assert q.pool == "A"
    assert q.topic == 1
    assert q.prompt_en.startswith("Which command entered on a switch")
    assert q.prompt_th and "คำสั่ง" in q.prompt_th
    assert [c["label"] for c in q.choices] == ["A", "B", "C", "D"]
    assert q.correct_labels == ["D"]
    assert q.needs_review is False
    assert q.image_filenames == []  # populated only by extract_table_images
    assert q.explanation and "Forward Time" in q.explanation


@needs_docx
def test_classify_table_handles_multi_answer(document):
    """At least one of Q23/Q30 from earlier scans should be a multi-answer."""
    multi_seen = False
    for i in range(60):  # search the first 60 questions
        q = classify_table(document.tables[i], source_index=i)
        if q is not None and len(q.correct_labels) >= 2:
            multi_seen = True
            break
    assert multi_seen, "expected to find at least one multi-answer question"


@needs_docx
def test_extract_table_images_writes_files(document, tmp_path):
    """Q2 is known to embed at least one image (per the probe scan)."""
    table = document.tables[1]  # 0-indexed → table 1 == Q#2
    files = extract_table_images(document, table, tmp_path, qid=2)
    assert files, "expected Q#2 to embed at least one image"
    for fname in files:
        path = tmp_path / fname
        assert path.exists() and path.stat().st_size > 0
        assert fname.startswith("Q-0002-")


# ---------- merged-cell choice reclamation tests (synthetic docx) ----------

def _set_cell_text(cell: _Cell, text: str, *, green: bool = False) -> None:
    """Replace *cell*'s text with a single paragraph + run, optionally green.

    Mirrors the shading shape the parser detects: ``<w:rPr><w:shd w:fill='e6f0e6'/></w:rPr>``.
    """
    # Clear existing paragraphs.
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    run = p.add_run(text)
    if green:
        rpr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), GREEN_FILL)
        rpr.append(shd)


def _merge_row(table: Table, ri: int) -> None:
    """Collapse *table*'s row *ri* into a single ``<w:tc>`` with ``gridSpan=2``.

    This mirrors the production-DOCX shape exactly: one underlying ``<w:tc>``
    spanning both columns, which python-docx then surfaces as ``row.cells``
    of length 2 — both entries returning identical text. Calling
    ``Cell.merge`` would concatenate the two cells' text instead, so we
    manipulate the XML directly.
    """
    row = table.rows[ri]
    tr = row._tr
    tcs = tr.findall(qn("w:tc"))
    if len(tcs) < 2:
        return
    tc0 = tcs[0]
    # Drop sibling tc's, keep only tc0.
    for tc in tcs[1:]:
        tr.remove(tc)
    # Mark tc0 as spanning both grid columns.
    tcpr = tc0.find(qn("w:tcPr"))
    if tcpr is None:
        tcpr = OxmlElement("w:tcPr")
        tc0.insert(0, tcpr)
    # Remove any pre-existing gridSpan, then add gridSpan=2.
    for existing in tcpr.findall(qn("w:gridSpan")):
        tcpr.remove(existing)
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), "2")
    tcpr.append(grid_span)


def _make_question_table_with_merged_e(
    *,
    qid: int = 999,
    topic: int = 1,
    pool: str = "A",
    prompt_en: str = "Synthetic prompt — pick the best answer.",
    choice_texts: tuple[str, str, str, str, str] = (
        "alpha", "bravo", "charlie", "delta", "echo",
    ),
    green_labels: tuple[str, ...] = ("E",),
    explanation: str | None = None,
) -> Table:
    """Build a 2-column question table whose 5th choice is a merged-cell row.

    Layout:
      row 0: header
      row 1: EN prompt (merged)
      row 2: choices A | B
      row 3: choices C | D
      row 4: choice E (merged single cell)
      row 5: optional explanation (merged)
    """
    doc = Document()
    n_rows = 5 + (1 if explanation else 0)
    table = doc.add_table(rows=n_rows, cols=2)

    # Row 0 — header.
    _set_cell_text(table.cell(0, 0), f"Question #{qid}")
    _set_cell_text(table.cell(0, 1), f"Topic {topic} - Exam Pool {pool}")

    # Row 1 — EN prompt, both cells identical so classifier sees a dup row.
    _set_cell_text(table.cell(1, 0), prompt_en)
    _set_cell_text(table.cell(1, 1), prompt_en)

    # Rows 2-3 — first four choices in 2-col layout.
    green_set = {f"A": choice_texts[0], "B": choice_texts[1],
                 "C": choice_texts[2], "D": choice_texts[3]}
    _set_cell_text(table.cell(2, 0), choice_texts[0], green="A" in green_labels)
    _set_cell_text(table.cell(2, 1), choice_texts[1], green="B" in green_labels)
    _set_cell_text(table.cell(3, 0), choice_texts[2], green="C" in green_labels)
    _set_cell_text(table.cell(3, 1), choice_texts[3], green="D" in green_labels)
    del green_set

    # Row 4 — choice E. Write the same text to both cells, then merge.
    _set_cell_text(table.cell(4, 0), choice_texts[4], green="E" in green_labels)
    _set_cell_text(table.cell(4, 1), choice_texts[4], green="E" in green_labels)
    _merge_row(table, 4)

    if explanation:
        _set_cell_text(table.cell(5, 0), explanation)
        _set_cell_text(table.cell(5, 1), explanation)

    return table


def test_extract_5_choices_when_5th_is_merged_row():
    """A 5-choice question whose 5th choice is a merged cell — choice E green.

    Reclamation rule: ``ri == last_choice_ri + 1`` AND green → reclaim as
    single-cell choice. Choice E must surface with the correct label.
    """
    table = _make_question_table_with_merged_e(
        qid=901,
        choice_texts=("alpha", "bravo", "charlie", "delta", "echo"),
        green_labels=("E",),
    )
    q = classify_table(table, source_index=0)

    assert q is not None
    assert q.id == 901
    assert [c["label"] for c in q.choices] == ["A", "B", "C", "D", "E"]
    assert [c["text"]  for c in q.choices] == [
        "alpha", "bravo", "charlie", "delta", "echo",
    ]
    assert q.correct_labels == ["E"]
    assert q.needs_review is False
    assert q.explanation is None


def test_choose_two_with_merged_e_choice():
    """A "Choose two" with green on A and the merged E — both must be correct."""
    table = _make_question_table_with_merged_e(
        qid=902,
        choice_texts=("alpha", "bravo", "charlie", "delta", "echo"),
        green_labels=("A", "E"),
    )
    q = classify_table(table, source_index=0)

    assert q is not None
    assert [c["label"] for c in q.choices] == ["A", "B", "C", "D", "E"]
    assert q.correct_labels == ["A", "E"]
    assert q.needs_review is False
    assert q.explanation is None


def test_explanation_after_merged_choice_e_stays_explanation():
    """Regression guard: when a merged choice E is followed by an explanation
    (the dual-merged-row case), reclamation must not eat the explanation.

    Rule: choice E is green → reclaimed; the following dup row is non-green
    and lives at ``last+2`` after reclaim, so it stays as explanation.
    """
    table = _make_question_table_with_merged_e(
        qid=903,
        choice_texts=("alpha", "bravo", "charlie", "delta", "echo"),
        green_labels=("E",),
        explanation="This is the explanation for the question.",
    )
    q = classify_table(table, source_index=0)

    assert q is not None
    assert [c["label"] for c in q.choices] == ["A", "B", "C", "D", "E"]
    assert q.correct_labels == ["E"]
    assert q.explanation == "This is the explanation for the question."
    assert q.needs_review is False


def test_short_english_explanation_not_misreclaimed_as_choice():
    """Regression guard: a NON-green merged row at last+1 with no further
    dup row is treated as an explanation, not a phantom choice E.

    This protects against false positives on 4-choice questions whose
    explanation happens to be short English prose.
    """
    doc = Document()
    table = doc.add_table(rows=5, cols=2)
    _set_cell_text(table.cell(0, 0), "Question #904")
    _set_cell_text(table.cell(0, 1), "Topic 1 - Exam Pool A")
    _set_cell_text(table.cell(1, 0), "Prompt EN.")
    _set_cell_text(table.cell(1, 1), "Prompt EN.")
    _set_cell_text(table.cell(2, 0), "alpha")
    _set_cell_text(table.cell(2, 1), "bravo", green=True)
    _set_cell_text(table.cell(3, 0), "charlie")
    _set_cell_text(table.cell(3, 1), "delta")
    # Merged non-green row at last+1: must be kept as explanation.
    _set_cell_text(table.cell(4, 0), "Brief English explanation.")
    _set_cell_text(table.cell(4, 1), "Brief English explanation.")
    _merge_row(table, 4)

    q = classify_table(table, source_index=0)
    assert q is not None
    assert [c["label"] for c in q.choices] == ["A", "B", "C", "D"]
    assert q.correct_labels == ["B"]
    assert q.explanation == "Brief English explanation."


# ---------- newline-preservation test for _cell_text ----------

def test_cell_text_preserves_paragraph_newlines():
    """Multi-paragraph cells must surface ``\\n`` between paragraphs so
    Cisco-config blocks and explanations keep their line structure.
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Clear the default empty paragraph python-docx inserts, then add three.
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    for line in ("line1", "line2", "line3"):
        cell.add_paragraph(line)

    assert _cell_text(cell) == "line1\nline2\nline3"


# ---------- multi-paragraph EN prompt join test ----------

def test_prompt_en_joins_multiple_english_dup_rows():
    """Q#435 pattern: a Cisco config block on one dup-row, the actual question
    on another. Both must surface in ``prompt_en`` separated by a blank line.
    Previously the parser used ``next()`` and silently dropped everything after
    the first English dup-row.
    """
    doc = Document()
    table = doc.add_table(rows=6, cols=2)
    _set_cell_text(table.cell(0, 0), "Question #905")
    _set_cell_text(table.cell(0, 1), "Topic 2 - Exam Pool B")
    # Row 1 — config block (English, first dup-row)
    config = "interface Gi1/0\nduplex full\nspeed 100"
    _set_cell_text(table.cell(1, 0), config)
    _set_cell_text(table.cell(1, 1), config)
    # Row 2 — actual question (English, second dup-row)
    question = "An engineer configures interface Gi1/0. Which action is necessary?"
    _set_cell_text(table.cell(2, 0), question)
    _set_cell_text(table.cell(2, 1), question)
    # Rows 3-4 — choices
    _set_cell_text(table.cell(3, 0), "alpha")
    _set_cell_text(table.cell(3, 1), "bravo", green=True)
    _set_cell_text(table.cell(4, 0), "charlie")
    _set_cell_text(table.cell(4, 1), "delta")
    # Row 5 — explanation (English, third dup-row but POST-choice, not joined)
    _set_cell_text(table.cell(5, 0), "Because LLDP-MED is required.")
    _set_cell_text(table.cell(5, 1), "Because LLDP-MED is required.")

    q = classify_table(table, source_index=0)
    assert q is not None
    assert q.prompt_en == f"{config}\n\n{question}"
    assert q.explanation == "Because LLDP-MED is required."
    assert [c["label"] for c in q.choices] == ["A", "B", "C", "D"]
    assert q.correct_labels == ["B"]


def test_prompt_th_joins_multiple_thai_dup_rows():
    """Same idea but for Thai prompts — translations occasionally span
    two paragraphs (context + question) in the source DOCX.
    """
    doc = Document()
    table = doc.add_table(rows=5, cols=2)
    _set_cell_text(table.cell(0, 0), "Question #906")
    _set_cell_text(table.cell(0, 1), "Topic 3 - Exam Pool C")
    th_context = "บริบทภาษาไทยส่วนต้น"
    th_question = "คำถามภาษาไทยส่วนที่สอง"
    _set_cell_text(table.cell(1, 0), th_context)
    _set_cell_text(table.cell(1, 1), th_context)
    _set_cell_text(table.cell(2, 0), th_question)
    _set_cell_text(table.cell(2, 1), th_question)
    _set_cell_text(table.cell(3, 0), "alpha", green=True)
    _set_cell_text(table.cell(3, 1), "bravo")
    _set_cell_text(table.cell(4, 0), "charlie")
    _set_cell_text(table.cell(4, 1), "delta")

    q = classify_table(table, source_index=0)
    assert q is not None
    assert q.prompt_th == f"{th_context}\n\n{th_question}"
    assert q.prompt_en == ""  # no English dup-rows
    assert q.correct_labels == ["A"]


def test_prompt_th_picks_mixed_english_thai_row():
    """Q#840 pattern: a Thai translation row that inlines many English
    technical terms (so ASCII letters > Thai chars). Must still bucket
    into ``prompt_th`` because it has Thai content, not into ``prompt_en``.
    """
    doc = Document()
    table = doc.add_table(rows=5, cols=2)
    _set_cell_text(table.cell(0, 0), "Question #907")
    _set_cell_text(table.cell(0, 1), "Topic 4 - Exam Pool D")
    en_prompt = "A network engineer must migrate a loopback interface to IPv6."
    # This Thai row is mostly English by ASCII-letter count but is still
    # the Thai translation. The fix buckets ANY Thai content as TH.
    th_mixed = ("วิศวกรเครือข่ายจะต้อง Migrate Router loopback interface "
                "ไปยัง IPv6 address space")
    _set_cell_text(table.cell(1, 0), en_prompt)
    _set_cell_text(table.cell(1, 1), en_prompt)
    _set_cell_text(table.cell(2, 0), th_mixed)
    _set_cell_text(table.cell(2, 1), th_mixed)
    _set_cell_text(table.cell(3, 0), "/64", green=True)
    _set_cell_text(table.cell(3, 1), "/96")
    _set_cell_text(table.cell(4, 0), "/124")
    _set_cell_text(table.cell(4, 1), "/128")

    q = classify_table(table, source_index=0)
    assert q is not None
    assert q.prompt_en == en_prompt
    assert q.prompt_th == th_mixed
    assert q.correct_labels == ["A"]
